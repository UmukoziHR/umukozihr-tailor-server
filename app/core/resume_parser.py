"""
Resume Parser Module
Extracts text from PDF/DOCX/TXT and uses Gemini 2.0 Pro to parse into ProfileV3 structure.
"""
import os
import io
import json
import logging
from typing import Optional, Tuple
from dotenv import load_dotenv
from google import genai
from google.genai.types import Schema, GenerateContentConfig

load_dotenv()
logger = logging.getLogger(__name__)

# Gemini 2.0 Pro system prompt for resume extraction
EXTRACTION_SYSTEM_PROMPT = """You are an expert resume parser. Your task is to extract structured data from the provided resume text.

RULES:
1. Extract ONLY information that is explicitly stated in the resume
2. Do NOT invent or hallucinate any data
3. If a field cannot be found, leave it as empty string "" or empty array []
4. Dates should be in YYYY-MM format (e.g., "2025-01" or "present")
5. Parse skills with appropriate levels (beginner, intermediate, expert)
6. Extract bullet points as they appear, preserving the original wording
7. For LinkedIn PDF exports, adapt to their specific format
8. LANGUAGES: Extract actual language names (English, French, Spanish, etc) - NEVER output "Language" as a placeholder
9. CERTIFICATIONS: Extract all certification names with issuer - include ALL found

Return ONLY valid JSON matching the exact schema provided."""

# ProfileV3 JSON Schema for Gemini structured output
PROFILE_V3_SCHEMA = Schema(
    type="OBJECT",
    required=["basics", "skills", "experience", "education", "projects", "certifications", "awards", "languages"],
    properties={
        "basics": Schema(type="OBJECT", properties={
            "full_name": Schema(type="STRING"),
            "headline": Schema(type="STRING"),
            "summary": Schema(type="STRING"),
            "location": Schema(type="STRING"),
            "email": Schema(type="STRING"),
            "phone": Schema(type="STRING"),
            "website": Schema(type="STRING"),
            "links": Schema(type="ARRAY", items=Schema(type="STRING")),
        }),
        "skills": Schema(type="ARRAY", items=Schema(type="OBJECT", properties={
            "name": Schema(type="STRING"),
            "level": Schema(type="STRING"),
            "keywords": Schema(type="ARRAY", items=Schema(type="STRING")),
        })),
        "experience": Schema(type="ARRAY", items=Schema(type="OBJECT", properties={
            "title": Schema(type="STRING"),
            "company": Schema(type="STRING"),
            "location": Schema(type="STRING"),
            "start": Schema(type="STRING"),
            "end": Schema(type="STRING"),
            "employment_type": Schema(type="STRING"),
            "bullets": Schema(type="ARRAY", items=Schema(type="STRING")),
        })),
        "education": Schema(type="ARRAY", items=Schema(type="OBJECT", properties={
            "school": Schema(type="STRING"),
            "degree": Schema(type="STRING"),
            "start": Schema(type="STRING"),
            "end": Schema(type="STRING"),
            "gpa": Schema(type="STRING"),
        })),
        "projects": Schema(type="ARRAY", items=Schema(type="OBJECT", properties={
            "name": Schema(type="STRING"),
            "url": Schema(type="STRING"),
            "stack": Schema(type="ARRAY", items=Schema(type="STRING")),
            "bullets": Schema(type="ARRAY", items=Schema(type="STRING")),
        })),
        "certifications": Schema(type="ARRAY", items=Schema(type="OBJECT", properties={
            "name": Schema(type="STRING"),
            "issuer": Schema(type="STRING"),
            "date": Schema(type="STRING"),
        })),
        "awards": Schema(type="ARRAY", items=Schema(type="OBJECT", properties={
            "name": Schema(type="STRING"),
            "by": Schema(type="STRING"),
            "date": Schema(type="STRING"),
        })),
        "languages": Schema(type="ARRAY", items=Schema(type="OBJECT", properties={
            "name": Schema(type="STRING"),
            "level": Schema(type="STRING"),
        })),
    },
)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using pdfplumber for better accuracy."""
    import pdfplumber
    
    text_parts = []
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        
        full_text = "\n\n".join(text_parts)
        logger.info(f"Extracted {len(full_text)} chars from PDF ({len(text_parts)} pages)")
        return full_text
    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        raise ValueError(f"Failed to extract text from PDF: {str(e)}")


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX file."""
    from docx import Document
    
    try:
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        
        full_text = "\n".join(paragraphs)
        logger.info(f"Extracted {len(full_text)} chars from DOCX")
        return full_text
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        raise ValueError(f"Failed to extract text from DOCX: {str(e)}")


def extract_text_from_txt(file_bytes: bytes) -> str:
    """Extract text from plain text file."""
    try:
        # Try UTF-8 first, then fall back to latin-1
        try:
            text = file_bytes.decode('utf-8')
        except UnicodeDecodeError:
            text = file_bytes.decode('latin-1')
        
        logger.info(f"Extracted {len(text)} chars from TXT")
        return text
    except Exception as e:
        logger.error(f"TXT extraction failed: {e}")
        raise ValueError(f"Failed to read text file: {str(e)}")


def extract_text(file_bytes: bytes, content_type: str, filename: str) -> str:
    """Extract text from uploaded file based on content type."""
    logger.info(f"Extracting text from {filename} (type: {content_type})")
    
    # Determine file type
    filename_lower = filename.lower()
    
    if content_type == "application/pdf" or filename_lower.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif content_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                          "application/msword"] or filename_lower.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    elif content_type == "text/plain" or filename_lower.endswith(".txt"):
        return extract_text_from_txt(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {content_type}. Supported: PDF, DOCX, TXT")


def parse_resume_with_llm(resume_text: str) -> dict:
    """Use Gemini 2.0 Pro to parse resume text into ProfileV3 structure."""
    logger.info(f"=== RESUME PARSING START === Text length: {len(resume_text)} chars")
    
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        logger.error("GEMINI_API_KEY not set")
        raise RuntimeError("GEMINI_API_KEY not set")
    
    try:
        client = genai.Client(api_key=api_key)
        
        user_prompt = f"""Parse the following resume and extract all information into the structured format.

RESUME TEXT:
{resume_text}

Extract all available information. For any field not found in the resume, use empty string "" or empty array [].
Return ONLY valid JSON matching the schema."""

        cfg = GenerateContentConfig(
            response_mime_type="application/json",
            response_schema=PROFILE_V3_SCHEMA,
            temperature=0.1,  # Low temperature for accurate extraction
            top_p=0.9,
            candidate_count=1,
            max_output_tokens=10000,
        )
        
        logger.info("Sending resume to Gemini 2.5-pro for parsing...")
        response = client.models.generate_content(
            model="gemini-2.5-pro",
            contents=[f"{EXTRACTION_SYSTEM_PROMPT}\n\n{user_prompt}"],
            config=cfg,
        )
        
        if not response.text:
            logger.error("Gemini returned empty response")
            raise RuntimeError("LLM returned empty response")
        
        logger.info(f"=== RESUME PARSING SUCCESS === Response length: {len(response.text)} chars")
        
        # Parse the JSON response
        parsed = json.loads(response.text)
        return parsed
        
    except json.JSONDecodeError as e:
        logger.error(f"Failed to parse LLM response as JSON: {e}")
        raise RuntimeError(f"Failed to parse resume extraction result: {str(e)}")
    except Exception as e:
        logger.error(f"Resume parsing failed: {e}", exc_info=True)
        raise


VALID_SKILL_LEVELS = {"beginner", "intermediate", "expert"}


def sanitize_extracted_profile(profile: dict) -> dict:
    """
    Sanitize extracted profile data to ensure it matches schema requirements.
    Fixes common extraction issues like empty skill levels and placeholder text.
    """
    # Sanitize skills - ensure valid level
    skills = profile.get("skills", [])
    for skill in skills:
        level = skill.get("level", "").lower().strip()
        if level not in VALID_SKILL_LEVELS:
            # Default to intermediate if level is missing or invalid
            skill["level"] = "intermediate"
            logger.debug(f"Skill '{skill.get('name', 'unknown')}' level defaulted to 'intermediate'")
    
    # Sanitize languages - remove placeholder entries like "Language"
    languages = profile.get("languages", [])
    sanitized_languages = []
    for lang in languages:
        name = lang.get("name", "").strip()
        # Skip if name is empty, just "Language", or other obvious placeholders
        if name and name.lower() not in ["language", "lang", ""]:
            sanitized_languages.append(lang)
        else:
            logger.debug(f"Filtered out placeholder language entry: {lang}")
    profile["languages"] = sanitized_languages
    
    # Sanitize certifications - remove placeholder entries
    certifications = profile.get("certifications", [])
    sanitized_certs = []
    for cert in certifications:
        name = cert.get("name", "").strip()
        # Skip if name is empty or placeholder
        if name and name.lower() not in ["certification", "cert", ""]:
            sanitized_certs.append(cert)
        else:
            logger.debug(f"Filtered out placeholder certification entry: {cert}")
    profile["certifications"] = sanitized_certs
    
    # Ensure basics has required fields with defaults
    if "basics" not in profile:
        profile["basics"] = {}
    basics = profile["basics"]
    basics.setdefault("full_name", "")
    basics.setdefault("headline", "")
    basics.setdefault("summary", "")
    basics.setdefault("location", "")
    basics.setdefault("email", "")
    basics.setdefault("phone", "")
    basics.setdefault("website", "")
    basics.setdefault("links", [])
    
    # Ensure all arrays exist
    profile.setdefault("skills", [])
    profile.setdefault("experience", [])
    profile.setdefault("education", [])
    profile.setdefault("projects", [])
    profile.setdefault("certifications", [])
    profile.setdefault("awards", [])
    profile.setdefault("languages", [])
    
    return profile


def calculate_extraction_confidence(profile: dict) -> Tuple[float, list]:
    """Calculate confidence score and list warnings for extracted profile."""
    warnings = []
    scores = []
    
    # Check basics
    basics = profile.get("basics", {})
    if basics.get("full_name"):
        scores.append(1.0)
    else:
        scores.append(0.0)
        warnings.append("Could not extract name")
    
    if basics.get("email"):
        scores.append(1.0)
    else:
        scores.append(0.0)
        warnings.append("Could not extract email")
    
    if basics.get("phone"):
        scores.append(1.0)
    else:
        scores.append(0.5)
        warnings.append("Could not extract phone number")
    
    # Check experience
    experience = profile.get("experience", [])
    if len(experience) > 0:
        scores.append(1.0)
    else:
        scores.append(0.3)
        warnings.append("No work experience found")
    
    # Check education
    education = profile.get("education", [])
    if len(education) > 0:
        scores.append(1.0)
    else:
        scores.append(0.5)
        warnings.append("No education found")
    
    # Check skills
    skills = profile.get("skills", [])
    if len(skills) > 0:
        scores.append(1.0)
    else:
        scores.append(0.5)
        warnings.append("No skills found")
    
    confidence = sum(scores) / len(scores) if scores else 0.0
    return round(confidence, 2), warnings


def parse_resume(file_bytes: bytes, content_type: str, filename: str) -> dict:
    """
    Main entry point: Extract text from file and parse with LLM.
    
    Returns:
        {
            "success": True,
            "profile": {...ProfileV3 structure...},
            "extraction_confidence": 0.85,
            "warnings": ["Could not extract phone number"],
            "message": "Resume parsed successfully"
        }
    """
    try:
        # Step 1: Extract text from file
        resume_text = extract_text(file_bytes, content_type, filename)
        
        if len(resume_text.strip()) < 50:
            return {
                "success": False,
                "profile": None,
                "extraction_confidence": 0.0,
                "warnings": ["File appears to be empty or unreadable"],
                "message": "Could not extract text from file"
            }
        
        # Step 2: Parse with LLM
        parsed_profile = parse_resume_with_llm(resume_text)
        
        # Step 3: Sanitize extracted data to ensure schema compliance
        parsed_profile = sanitize_extracted_profile(parsed_profile)
        logger.info("Profile sanitized for schema compliance")
        
        # Step 4: Calculate confidence
        confidence, warnings = calculate_extraction_confidence(parsed_profile)
        
        return {
            "success": True,
            "profile": parsed_profile,
            "extraction_confidence": confidence,
            "warnings": warnings,
            "message": "Resume parsed successfully"
        }
        
    except ValueError as e:
        logger.error(f"Resume parsing error: {e}")
        return {
            "success": False,
            "profile": None,
            "extraction_confidence": 0.0,
            "warnings": [str(e)],
            "message": str(e)
        }
    except Exception as e:
        logger.error(f"Unexpected error in parse_resume: {e}", exc_info=True)
        return {
            "success": False,
            "profile": None,
            "extraction_confidence": 0.0,
            "warnings": ["An unexpected error occurred"],
            "message": f"Failed to parse resume: {str(e)}"
        }
