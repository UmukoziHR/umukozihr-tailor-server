"""
DOCX Document Generator for UmukoziHR Resume Tailor
Modern, ATS-friendly design matching PDF templates.
Supports regional formats: US, EU, Global.
"""

import os
import logging
from calendar import month_name
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)

# Use ARTIFACTS_DIR env var if set, otherwise fallback to local path
ART_DIR = os.environ.get("ARTIFACTS_DIR", os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", "..", "artifacts")))
os.makedirs(ART_DIR, exist_ok=True)

# === COLORS - Matching PDF templates ===
HEADER_COLOR = RGBColor(44, 62, 80)    # Dark slate blue
ACCENT_COLOR = RGBColor(52, 73, 94)    # Slightly lighter
SUBTLE_COLOR = RGBColor(100, 100, 100) # Gray for dates


def format_date_human(date_str: str) -> str:
    """Convert YYYY-MM to human readable format like 'June 2025'."""
    if not date_str or not isinstance(date_str, str):
        return date_str or ""
    
    date_str = date_str.strip()
    
    if date_str.lower() == 'present':
        return 'Present'
    
    if any(c.isalpha() for c in date_str):
        return date_str
    
    try:
        if '-' in date_str:
            parts = date_str.split('-')
            if len(parts) >= 2:
                year = parts[0]
                month_num = int(parts[1])
                if 1 <= month_num <= 12:
                    return f"{month_name[month_num]} {year}"
        if date_str.isdigit() and len(date_str) == 4:
            return date_str
    except (ValueError, IndexError):
        pass
    
    return date_str


def set_document_margins(doc, region: str = "GL"):
    """Set document margins based on region"""
    margins = {
        "US": 0.5,   # Tight for 1-page fit
        "EU": 0.7,   # Standard European
        "GL": 0.6    # Balanced global
    }
    margin = margins.get(region, 0.6)
    
    for section in doc.sections:
        section.top_margin = Inches(margin)
        section.bottom_margin = Inches(margin)
        section.left_margin = Inches(margin)
        section.right_margin = Inches(margin)


def add_horizontal_line(paragraph, color=HEADER_COLOR):
    """Add a colored horizontal line after a paragraph"""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '1')
    # Convert RGBColor to hex string
    color_hex = f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_section_header(doc, text: str, region: str = "GL"):
    """Add a styled section header matching PDF design"""
    header = doc.add_paragraph()
    run = header.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = HEADER_COLOR
    
    # Small caps effect through letter spacing
    header.paragraph_format.space_before = Pt(14 if region == "EU" else 12)
    header.paragraph_format.space_after = Pt(2)
    
    # Add underline
    add_horizontal_line(header, HEADER_COLOR)
    
    return header


def create_resume_docx(profile: dict, resume_out: dict, job: dict, out_path: str, region: str = "GL") -> str:
    """
    Create a professional resume DOCX document.
    
    Args:
        profile: User profile data (name, contacts, etc.)
        resume_out: LLM-generated resume content
        job: Job details (company, title, region)
        out_path: Output file path (without extension)
        region: US/EU/GL for regional formatting
    
    Returns:
        Path to the generated DOCX file
    """
    doc = Document()
    set_document_margins(doc, region)
    
    # === HEADER: Name ===
    name_para = doc.add_paragraph()
    name_run = name_para.add_run(profile.get('name', 'Your Name'))
    name_run.bold = True
    name_run.font.size = Pt(20 if region == "EU" else 18)
    name_run.font.color.rgb = HEADER_COLOR
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_after = Pt(6)
    
    # === Contact Info ===
    contacts = profile.get('contacts', {})
    contact_parts = []
    if contacts.get('email'):
        contact_parts.append(f"📧 {contacts['email']}")
    if contacts.get('phone'):
        contact_parts.append(f"📱 {contacts['phone']}")
    if contacts.get('location'):
        contact_parts.append(f"📍 {contacts['location']}")
    
    if contact_parts:
        contact_para = doc.add_paragraph()
        contact_para.add_run('  •  '.join(contact_parts))
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_para.paragraph_format.space_after = Pt(4)
    
    # Links
    links = contacts.get('links', [])
    if links:
        links_para = doc.add_paragraph()
        links_para.add_run('  |  '.join(links[:3]))
        links_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        links_para.paragraph_format.space_after = Pt(8)
    
    # Separator line
    hr_para = doc.add_paragraph()
    add_horizontal_line(hr_para, HEADER_COLOR)
    
    # === PROFESSIONAL SUMMARY ===
    summary = resume_out.get('summary', '')
    if summary:
        add_section_header(doc, "Professional Summary" if region == "EU" else "Summary", region)
        summary_para = doc.add_paragraph()
        summary_para.add_run(summary)
        summary_para.paragraph_format.space_after = Pt(8)
    
    # === EXPERIENCE ===
    experience = resume_out.get('experience', [])
    if experience:
        add_section_header(doc, "Professional Experience" if region == "EU" else "Experience", region)
        
        for i, exp in enumerate(experience):
            # Title and Company
            role_para = doc.add_paragraph()
            title_run = role_para.add_run(exp.get('title', 'Role'))
            title_run.bold = True
            role_para.add_run('  |  ')
            company_run = role_para.add_run(exp.get('company', 'Company'))
            company_run.italic = True
            role_para.paragraph_format.space_before = Pt(10 if i > 0 else 6)
            role_para.paragraph_format.space_after = Pt(2)
            
            # Dates (styled gray)
            dates_para = doc.add_paragraph()
            start = format_date_human(exp.get('start', ''))
            end = format_date_human(exp.get('end', 'Present'))
            date_run = dates_para.add_run(f"{start} – {end}")
            date_run.font.color.rgb = SUBTLE_COLOR
            date_run.font.size = Pt(10)
            dates_para.paragraph_format.space_after = Pt(4)
            
            # Bullets
            for bullet in exp.get('bullets', []):
                bullet_para = doc.add_paragraph(style='List Bullet')
                bullet_para.add_run(bullet)
                bullet_para.paragraph_format.space_after = Pt(2)
                bullet_para.paragraph_format.left_indent = Inches(0.25)
    
    # === EDUCATION ===
    education = resume_out.get('education', [])
    if education:
        add_section_header(doc, "Education", region)
        
        for edu in education:
            edu_para = doc.add_paragraph()
            degree_run = edu_para.add_run(edu.get('degree', 'Degree'))
            degree_run.bold = True
            edu_para.add_run(f"  –  {edu.get('school', 'University')}")
            
            period = edu.get('period', '')
            if period:
                edu_para.add_run(f"  ({period})")
            
            edu_para.paragraph_format.space_after = Pt(6)
    
    # === SKILLS ===
    skills = resume_out.get('skills_line', resume_out.get('skills', []))
    if skills:
        add_section_header(doc, "Core Skills" if region == "GL" else "Skills", region)
        skills_para = doc.add_paragraph()
        skills_para.add_run('  •  '.join(skills))
        skills_para.paragraph_format.space_after = Pt(8)
    
    # === PROJECTS ===
    projects = resume_out.get('projects', [])
    if projects:
        add_section_header(doc, "Projects", region)
        
        for proj in projects:
            proj_para = doc.add_paragraph()
            name_run = proj_para.add_run(proj.get('name', 'Project'))
            name_run.bold = True
            
            stack = proj.get('stack', [])
            if stack:
                if isinstance(stack, list):
                    stack_str = ', '.join(stack)
                else:
                    stack_str = str(stack)
                proj_para.add_run(f"  |  {stack_str}")
            
            proj_para.paragraph_format.space_after = Pt(4)
            
            for bullet in proj.get('bullets', []):
                bullet_para = doc.add_paragraph(style='List Bullet')
                bullet_para.add_run(bullet)
                bullet_para.paragraph_format.space_after = Pt(2)
    
    # === CERTIFICATIONS ===
    certifications = resume_out.get('certifications', [])
    if certifications:
        add_section_header(doc, "Certifications", region)
        
        for cert in certifications:
            cert_para = doc.add_paragraph()
            cert_para.add_run(f"• {cert.get('name', 'Certification')}")
            issuer = cert.get('issuer', '')
            date = cert.get('date', '')
            if issuer or date:
                cert_para.add_run(f"  –  {issuer}" + (f" ({date})" if date else ""))
            cert_para.paragraph_format.space_after = Pt(3)
    
    # === AWARDS ===
    awards = resume_out.get('awards', [])
    if awards:
        add_section_header(doc, "Awards & Achievements" if region == "EU" else "Awards", region)
        
        for award in awards:
            award_para = doc.add_paragraph()
            award_para.add_run(f"• {award.get('name', 'Award')}")
            by = award.get('by', '')
            date = award.get('date', '')
            if by or date:
                award_para.add_run(f"  –  {by}" + (f" ({date})" if date else ""))
            award_para.paragraph_format.space_after = Pt(3)
    
    # === LANGUAGES ===
    languages = resume_out.get('languages', [])
    if languages:
        add_section_header(doc, "Languages", region)
        lang_para = doc.add_paragraph()
        lang_texts = [f"{l.get('name', '')} ({l.get('level', 'Fluent')})" for l in languages if l.get('name')]
        lang_para.add_run('  •  '.join(lang_texts))
    
    # EU style: Add references note
    if region == "EU":
        doc.add_paragraph()
        ref_para = doc.add_paragraph()
        ref_run = ref_para.add_run("References available upon request")
        ref_run.font.size = Pt(9)
        ref_run.font.color.rgb = SUBTLE_COLOR
        ref_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save
    docx_path = f"{out_path}_resume.docx"
    doc.save(docx_path)
    logger.info(f"Resume DOCX created: {docx_path}")
    
    return docx_path


def create_cover_letter_docx(profile: dict, cover_letter_out: dict, job: dict, out_path: str, region: str = "GL") -> str:
    """
    Create a professional cover letter DOCX document.
    
    Args:
        profile: User profile data
        cover_letter_out: LLM-generated cover letter content
        job: Job details
        out_path: Output file path
        region: US/EU/GL for regional formatting
    
    Returns:
        Path to the generated DOCX file
    """
    doc = Document()
    
    # Regional margins
    margins = {"US": 1.0, "EU": 1.0, "GL": 1.0}
    margin = margins.get(region, 1.0)
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(margin)
        section.right_margin = Inches(margin)
    
    contacts = profile.get('contacts', {})
    
    # === SENDER HEADER ===
    name_para = doc.add_paragraph()
    name_run = name_para.add_run(profile.get('name', 'Your Name'))
    name_run.bold = True
    name_run.font.size = Pt(14)
    name_run.font.color.rgb = HEADER_COLOR
    
    # Contact details
    if contacts.get('email'):
        email_para = doc.add_paragraph()
        email_para.add_run(f"📧 {contacts['email']}")
        email_para.paragraph_format.space_after = Pt(0)
    
    if contacts.get('phone'):
        phone_para = doc.add_paragraph()
        phone_para.add_run(f"📱 {contacts['phone']}")
        phone_para.paragraph_format.space_after = Pt(0)
    
    if contacts.get('location'):
        loc_para = doc.add_paragraph()
        loc_para.add_run(f"📍 {contacts['location']}")
    
    # Separator
    hr_para = doc.add_paragraph()
    add_horizontal_line(hr_para, HEADER_COLOR)
    
    # === DATE ===
    doc.add_paragraph()
    date_para = doc.add_paragraph()
    date_para.add_run(datetime.now().strftime("%B %d, %Y"))
    date_para.paragraph_format.space_after = Pt(12)
    
    # === RECIPIENT ===
    recipient_para = doc.add_paragraph()
    recipient_para.add_run("Hiring Manager")
    recipient_para.paragraph_format.space_after = Pt(0)
    
    company_para = doc.add_paragraph()
    company_para.add_run(job.get('company', 'Company Name'))
    company_para.paragraph_format.space_after = Pt(12)
    
    # === SUBJECT ===
    subject_para = doc.add_paragraph()
    subject_text = f"Re: Application for {job.get('title', 'Position')}"
    if region == "EU":
        subject_text = f"Application for the position of {job.get('title', 'Position')}"
    subject_run = subject_para.add_run(subject_text)
    subject_run.bold = True
    subject_para.paragraph_format.space_after = Pt(12)
    
    # === GREETING ===
    greeting = cover_letter_out.get('address', 'Dear Hiring Manager,')
    greeting_para = doc.add_paragraph()
    greeting_para.add_run(greeting)
    greeting_para.paragraph_format.space_after = Pt(10)
    
    # === BODY ===
    # Introduction
    intro = cover_letter_out.get('intro', '')
    if intro:
        intro_para = doc.add_paragraph()
        intro_para.add_run(intro)
        intro_para.paragraph_format.space_after = Pt(10)
        intro_para.paragraph_format.line_spacing = 1.15
    
    # Why you
    why_you = cover_letter_out.get('why_you', '')
    if why_you:
        why_para = doc.add_paragraph()
        why_para.add_run(why_you)
        why_para.paragraph_format.space_after = Pt(10)
        why_para.paragraph_format.line_spacing = 1.15
    
    # Evidence bullets
    evidence = cover_letter_out.get('evidence', [])
    if evidence:
        if region == "EU":
            intro_text = doc.add_paragraph()
            intro_text.add_run("I would like to highlight the following relevant achievements:")
            intro_text.paragraph_format.space_after = Pt(4)
        
        for point in evidence:
            bullet_para = doc.add_paragraph(style='List Bullet')
            bullet_para.add_run(point)
            bullet_para.paragraph_format.space_after = Pt(4)
            bullet_para.paragraph_format.left_indent = Inches(0.25)
        
        doc.add_paragraph()
    
    # Why them
    why_them = cover_letter_out.get('why_them', '')
    if why_them:
        why_them_para = doc.add_paragraph()
        why_them_para.add_run(why_them)
        why_them_para.paragraph_format.space_after = Pt(10)
        why_them_para.paragraph_format.line_spacing = 1.15
    
    # Closing
    close_text = cover_letter_out.get('close', '')
    if close_text:
        close_para = doc.add_paragraph()
        close_para.add_run(close_text)
        close_para.paragraph_format.space_after = Pt(16)
        close_para.paragraph_format.line_spacing = 1.15
    
    # === SIGNATURE ===
    closing_salutation = {
        "US": "Best regards,",
        "EU": "Yours sincerely,",
        "GL": "Sincerely,"
    }
    closing_para = doc.add_paragraph()
    closing_para.add_run(closing_salutation.get(region, "Sincerely,"))
    closing_para.paragraph_format.space_after = Pt(24)
    
    sig_para = doc.add_paragraph()
    sig_run = sig_para.add_run(profile.get('name', 'Your Name'))
    sig_run.bold = True
    
    # Save
    docx_path = f"{out_path}_cover.docx"
    doc.save(docx_path)
    logger.info(f"Cover letter DOCX created: {docx_path}")
    
    return docx_path


def render_docx(resume_ctx: dict, cl_ctx: dict, out_base: str, region: str = "GL") -> tuple:
    """
    Render both resume and cover letter as DOCX files.
    
    Args:
        resume_ctx: Context dict with 'profile', 'out' (resume), and 'job' keys
        cl_ctx: Context dict with 'profile', 'out' (cover letter), and 'job' keys
        out_base: Base filename (without extension)
        region: US/EU/GL for regional formatting
    
    Returns:
        Tuple of (resume_docx_path, cover_letter_docx_path)
    """
    out_path = os.path.join(ART_DIR, out_base)
    
    try:
        resume_docx = create_resume_docx(
            profile=resume_ctx['profile'],
            resume_out=resume_ctx['out'],
            job=resume_ctx['job'],
            out_path=out_path,
            region=region
        )
    except Exception as e:
        logger.error(f"Failed to create resume DOCX: {e}")
        resume_docx = None
    
    try:
        cover_docx = create_cover_letter_docx(
            profile=cl_ctx['profile'],
            cover_letter_out=cl_ctx['out'],
            job=cl_ctx['job'],
            out_path=out_path,
            region=region
        )
    except Exception as e:
        logger.error(f"Failed to create cover letter DOCX: {e}")
        cover_docx = None
    
    return resume_docx, cover_docx
